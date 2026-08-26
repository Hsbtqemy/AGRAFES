"""Soft line breaks must open a new line in the numbered-lines importers.

A DOCX ``<w:br/>`` / ODT ``text:line-break`` reaches the parser as a newline
*inside* a paragraph. Keying units on ``<w:p>`` alone collapsed a whole
break-separated document into a single unit — the "blob" case described in
``docs/DESIGN_R2_3_blob_two_grain.md`` §0, and the shape of the 15 single-unit
documents found in the working corpus (bitexts exported as one Word paragraph
carrying ~900-1300 breaks).

Two failure modes were observed and are both covered here:
- the paragraph starts with ``[n]`` → ``re.DOTALL`` let ``.+`` swallow the whole
  file into ONE ``line`` unit;
- it starts with anything else → no match at all → ONE ``structure`` unit, not
  even indexed in FTS.
"""

from __future__ import annotations

import io
import zipfile
from pathlib import Path
from types import SimpleNamespace

import pytest

docx = pytest.importorskip("docx")

from multicorpus_engine.importers.docx_numbered_lines import (  # noqa: E402
    parse_docx_numbered_lines,
)
from multicorpus_engine.importers.odt_common import (  # noqa: E402
    read_odt_paragraph_rich_lines,
)
from multicorpus_engine.importers.odt_numbered_lines import (  # noqa: E402
    parse_odt_numbered_lines,
)
from multicorpus_engine.importers.rich_text import para_to_rich_lines  # noqa: E402


# ─── helpers ────────────────────────────────────────────────────────────────


def _docx_with_paragraphs(tmp_path: Path, *texts: str) -> Path:
    """Write a DOCX whose paragraphs hold *texts* — ``\n`` becomes ``<w:br/>``."""
    doc = docx.Document()
    for text in texts:
        doc.add_paragraph(text)
    path = tmp_path / "soft_breaks.docx"
    doc.save(str(path))
    return path


def _odt_with_raw_paragraphs(tmp_path: Path, *raw_paragraph_xml: str) -> Path:
    """Write an ODT whose ``text:p`` bodies are given verbatim (markup allowed)."""
    text_ns = "urn:oasis:names:tc:opendocument:xmlns:text:1.0"
    office_ns = "urn:oasis:names:tc:opendocument:xmlns:office:1.0"
    inner = "".join(f"<text:p>{body}</text:p>" for body in raw_paragraph_xml)
    content_xml = (
        '<?xml version="1.0" encoding="UTF-8"?>'
        f'<office:document-content xmlns:office="{office_ns}" '
        f'xmlns:text="{text_ns}" office:version="1.3">'
        f"<office:body><office:text>{inner}</office:text></office:body>"
        "</office:document-content>"
    )
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr(
            "mimetype",
            "application/vnd.oasis.opendocument.text",
            compress_type=zipfile.ZIP_STORED,
        )
        zf.writestr("content.xml", content_xml.encode("utf-8"))
    path = tmp_path / "soft_breaks.odt"
    path.write_bytes(buf.getvalue())
    return path


# ─── DOCX ───────────────────────────────────────────────────────────────────


def test_docx_soft_breaks_open_new_numbered_lines(tmp_path: Path) -> None:
    """One paragraph, three ``[n]`` lines separated by soft breaks → three units."""
    path = _docx_with_paragraphs(tmp_path, "[1] alpha\n[2] beta\n[3] gamma")

    units = parse_docx_numbered_lines(path).units

    assert [(u.unit_type, u.external_id, u.text_raw) for u in units] == [
        ("line", 1, "alpha"),
        ("line", 2, "beta"),
        ("line", 3, "gamma"),
    ]
    assert [u.n for u in units] == [1, 2, 3]


def test_docx_soft_breaks_below_an_unnumbered_head(tmp_path: Path) -> None:
    """A title line above the numbering must not swallow the body into one structure unit."""
    path = _docx_with_paragraphs(tmp_path, "Titre du document\n[1] alpha\n[2] beta")

    units = parse_docx_numbered_lines(path).units

    assert [(u.unit_type, u.external_id) for u in units] == [
        ("structure", None),
        ("line", 1),
        ("line", 2),
    ]
    assert units[0].text_raw == "Titre du document"


def test_docx_blank_lines_between_breaks_are_dropped(tmp_path: Path) -> None:
    """Consecutive breaks produce empty lines — skipped like blank paragraphs."""
    path = _docx_with_paragraphs(tmp_path, "[1] alpha\n\n\n[2] beta")

    units = parse_docx_numbered_lines(path).units

    assert [(u.n, u.external_id) for u in units] == [(1, 1), (2, 2)]


def test_docx_paragraph_without_break_is_unchanged(tmp_path: Path) -> None:
    """Regression: the common case (one line per ``<w:p>``) keeps its exact output."""
    path = _docx_with_paragraphs(tmp_path, "[1] alpha", "intertitre", "[2] beta")

    units = parse_docx_numbered_lines(path).units

    assert [(u.n, u.unit_type, u.external_id, u.text_raw) for u in units] == [
        (1, "line", 1, "alpha"),
        (2, "structure", None, "intertitre"),
        (3, "line", 2, "beta"),
    ]


def test_docx_style_spanning_a_break_stays_balanced(tmp_path: Path) -> None:
    """A styled run straddling a break yields one closed ``<hi>`` per line."""
    doc = docx.Document()
    para = doc.add_paragraph()
    run = para.add_run("alpha\nbeta")
    run.italic = True
    path = tmp_path / "styled.docx"
    doc.save(str(path))

    lines = para_to_rich_lines(docx.Document(str(path)).paragraphs[0])

    assert lines == ['<hi rend="italic">alpha</hi>', '<hi rend="italic">beta</hi>']


def test_docx_marker_keeps_sep_count_meta(tmp_path: Path) -> None:
    """The ¤ separator (ADR-002) is still counted per line, not per paragraph."""
    path = _docx_with_paragraphs(tmp_path, "[1] a ¤ b\n[2] c")

    units = parse_docx_numbered_lines(path).units

    assert units[0].meta_json == '{"sep_count": 1}'
    assert units[1].meta_json is None


# ─── ODT ────────────────────────────────────────────────────────────────────


def test_odt_soft_breaks_open_new_numbered_lines(tmp_path: Path) -> None:
    path = _odt_with_raw_paragraphs(
        tmp_path,
        "[1] alpha<text:line-break/>[2] beta<text:line-break/>[3] gamma",
    )

    units = parse_odt_numbered_lines(path).units

    assert [(u.unit_type, u.external_id, u.text_raw) for u in units] == [
        ("line", 1, "alpha"),
        ("line", 2, "beta"),
        ("line", 3, "gamma"),
    ]


def test_odt_paragraph_reader_is_left_alone(tmp_path: Path) -> None:
    """``odt_paragraphs`` semantics (one unit per ``text:p``) must not shift:
    only the numbered-lines importer reads the split view."""
    path = _odt_with_raw_paragraphs(tmp_path, "alpha<text:line-break/>beta")

    assert read_odt_paragraph_rich_lines(path) == [("alpha\nbeta", None)]


def test_docx_column_extraction_counts_lines_not_paragraphs(tmp_path: Path) -> None:
    """Column walk: a cell holding soft breaks yields one unit per line, and the
    unnumbered-ratio counters weigh lines — not the single ``<w:p>`` holding them."""
    doc = docx.Document()
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).paragraphs[0].text = "[1] source une\n[2] source deux"
    table.cell(0, 1).paragraphs[0].text = "[1] cible une\n[2] cible deux"
    path = tmp_path / "colonnes.docx"
    doc.save(str(path))

    parsed = parse_docx_numbered_lines(path, column_index=2)

    assert [(u.external_id, u.text_raw) for u in parsed.units] == [
        (1, "cible une"),
        (2, "cible deux"),
    ]
    assert parsed.stats["col_paragraphs_total"] == 2
    assert parsed.stats["col_paragraphs_line"] == 2


def test_docx_page_break_does_not_split_a_line(tmp_path: Path) -> None:
    """Only *line* breaks split. python-docx ignores page/column breaks in
    ``Run.text``, so a page break must leave the line whole."""
    from docx.enum.text import WD_BREAK

    doc = docx.Document()
    para = doc.add_paragraph()
    run = para.add_run("[1] avant et ")
    run.add_break(WD_BREAK.PAGE)
    para.add_run("après")
    path = tmp_path / "saut_de_page.docx"
    doc.save(str(path))

    units = parse_docx_numbered_lines(path).units

    assert [(u.unit_type, u.external_id, u.text_raw) for u in units] == [
        ("line", 1, "avant et après"),
    ]


def test_carriage_returns_are_treated_as_line_breaks() -> None:
    """Defensive: XML normalises CRLF away before we see it, but the splitter
    must not leave a stray ``\r`` inside a line if one ever reaches it."""
    para = SimpleNamespace(
        runs=[
            SimpleNamespace(
                text="alpha\r\nbeta\rgamma",
                bold=None, italic=None, underline=None,
                font=SimpleNamespace(strike=None, superscript=None, subscript=None),
            )
        ]
    )

    assert para_to_rich_lines(para) == ["alpha", "beta", "gamma"]
