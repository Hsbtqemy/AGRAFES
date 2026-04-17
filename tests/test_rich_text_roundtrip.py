"""RICH-7 — Round-trip integration tests: import with inline markup → export TEI.

Tests verify that:
- Italic/bold/underline runs in DOCX are preserved in text_raw as <hi rend="…">
- text_norm is always plain text (no <hi> tags)
- The TEI exporter propagates <hi> tags as XML sub-elements
"""

from __future__ import annotations

import io
import sqlite3
import xml.etree.ElementTree as ET
from pathlib import Path

import pytest


_TEI_NS = "http://www.tei-c.org/ns/1.0"


def _ns(tag: str) -> str:
    return f"{{{_TEI_NS}}}{tag}"


# ─── DOCX helpers ─────────────────────────────────────────────────────────────

def _make_docx_with_runs(
    paragraphs: list[list[tuple[str, dict[str, bool]]]],
) -> bytes:
    """Build a DOCX in memory.

    paragraphs: list of paragraphs, each paragraph is a list of (text, style_kwargs)
    where style_kwargs can contain italic=True, bold=True, underline=True, etc.

    Example:
        [
            [("plain ", {}), ("italic text", {"italic": True}), (" end", {})],
        ]
    """
    import docx

    doc = docx.Document()
    for para_runs in paragraphs:
        para = doc.add_paragraph()
        for text, styles in para_runs:
            run = para.add_run(text)
            for attr, val in styles.items():
                setattr(run, attr, val)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_docx_numbered_with_runs(
    entries: list[tuple[int, list[tuple[str, dict[str, bool]]]]],
) -> bytes:
    """Build a numbered-lines DOCX with styled runs.

    entries: list of (external_id, [(text, style_kwargs), ...])
    The [n] prefix is always added as plain text, then the styled runs follow.
    """
    import docx

    doc = docx.Document()
    for ext_id, runs in entries:
        para = doc.add_paragraph()
        para.add_run(f"[{ext_id}] ")
        for text, styles in runs:
            run = para.add_run(text)
            for attr, val in styles.items():
                setattr(run, attr, val)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ─── Tests: docx_paragraphs ───────────────────────────────────────────────────

class TestDocxParagraphsRichImport:
    def test_italic_preserved_in_text_raw(
        self, db_conn: sqlite3.Connection, tmp_path: Path
    ) -> None:
        """Italic run must appear as <hi rend="italic"> in text_raw."""
        from multicorpus_engine.importers.docx_paragraphs import import_docx_paragraphs

        data = _make_docx_with_runs([
            [("un ", {}), ("titre", {"italic": True}), (" important", {})],
        ])
        p = tmp_path / "italic.docx"
        p.write_bytes(data)

        report = import_docx_paragraphs(db_conn, path=p, language="fr")

        row = db_conn.execute(
            "SELECT text_raw, text_norm FROM units WHERE doc_id = ?", (report.doc_id,)
        ).fetchone()
        assert row is not None
        assert '<hi rend="italic">' in row["text_raw"], f"text_raw={row['text_raw']!r}"
        assert "titre" in row["text_raw"]
        assert "<hi" not in row["text_norm"], f"text_norm should be plain: {row['text_norm']!r}"
        assert "titre" in row["text_norm"]

    def test_bold_preserved_in_text_raw(
        self, db_conn: sqlite3.Connection, tmp_path: Path
    ) -> None:
        data = _make_docx_with_runs([
            [("introduction ", {}), ("importante", {"bold": True})],
        ])
        p = tmp_path / "bold.docx"
        p.write_bytes(data)

        from multicorpus_engine.importers.docx_paragraphs import import_docx_paragraphs
        report = import_docx_paragraphs(db_conn, path=p, language="fr")

        row = db_conn.execute(
            "SELECT text_raw, text_norm FROM units WHERE doc_id = ?", (report.doc_id,)
        ).fetchone()
        assert '<hi rend="bold">' in row["text_raw"]
        assert "<hi" not in row["text_norm"]

    def test_bold_italic_combined(
        self, db_conn: sqlite3.Connection, tmp_path: Path
    ) -> None:
        data = _make_docx_with_runs([
            [("texte ", {}), ("fort et penché", {"bold": True, "italic": True})],
        ])
        p = tmp_path / "bold_italic.docx"
        p.write_bytes(data)

        from multicorpus_engine.importers.docx_paragraphs import import_docx_paragraphs
        report = import_docx_paragraphs(db_conn, path=p, language="fr")

        row = db_conn.execute(
            "SELECT text_raw FROM units WHERE doc_id = ?", (report.doc_id,)
        ).fetchone()
        # bold and italic → rend="bold italic" (alphabetical)
        assert 'rend="bold italic"' in row["text_raw"], f"Got: {row['text_raw']!r}"

    def test_plain_paragraph_unchanged(
        self, db_conn: sqlite3.Connection, tmp_path: Path
    ) -> None:
        """A paragraph with no styling must not contain any <hi> tags."""
        data = _make_docx_with_runs([
            [("Texte ordinaire sans mise en forme.", {})],
        ])
        p = tmp_path / "plain.docx"
        p.write_bytes(data)

        from multicorpus_engine.importers.docx_paragraphs import import_docx_paragraphs
        report = import_docx_paragraphs(db_conn, path=p, language="fr")

        row = db_conn.execute(
            "SELECT text_raw, text_norm FROM units WHERE doc_id = ?", (report.doc_id,)
        ).fetchone()
        assert "<hi" not in row["text_raw"]
        assert row["text_raw"] == row["text_norm"]

    def test_text_norm_is_always_plain(
        self, db_conn: sqlite3.Connection, tmp_path: Path
    ) -> None:
        """text_norm must never contain <hi> markup regardless of input."""
        data = _make_docx_with_runs([
            [("a", {"italic": True}), ("b", {"bold": True}), ("c", {})],
        ])
        p = tmp_path / "mixed.docx"
        p.write_bytes(data)

        from multicorpus_engine.importers.docx_paragraphs import import_docx_paragraphs
        report = import_docx_paragraphs(db_conn, path=p, language="fr")

        rows = db_conn.execute(
            "SELECT text_norm FROM units WHERE doc_id = ?", (report.doc_id,)
        ).fetchall()
        for row in rows:
            assert "<hi" not in (row["text_norm"] or ""), f"text_norm has markup: {row['text_norm']!r}"


# ─── Tests: docx_numbered_lines ───────────────────────────────────────────────

class TestDocxNumberedLinesRichImport:
    def test_italic_in_content_preserved(
        self, db_conn: sqlite3.Connection, tmp_path: Path
    ) -> None:
        """[n] prefix is plain; styled content after prefix must be preserved."""
        data = _make_docx_numbered_with_runs([
            (1, [("titre ", {}), ("en italique", {"italic": True})]),
            (2, [("texte ordinaire", {})]),
        ])
        p = tmp_path / "numbered_italic.docx"
        p.write_bytes(data)

        from multicorpus_engine.importers.docx_numbered_lines import import_docx_numbered_lines
        report = import_docx_numbered_lines(db_conn, path=p, language="fr")

        rows = {
            row["external_id"]: row
            for row in db_conn.execute(
                "SELECT external_id, text_raw, text_norm FROM units WHERE doc_id = ? AND unit_type = 'line'",
                (report.doc_id,),
            ).fetchall()
        }
        assert '<hi rend="italic">' in rows[1]["text_raw"], f"Got: {rows[1]['text_raw']!r}"
        assert "<hi" not in rows[1]["text_norm"]
        assert "<hi" not in rows[2]["text_raw"]

    def test_numbered_prefix_not_in_text_raw(
        self, db_conn: sqlite3.Connection, tmp_path: Path
    ) -> None:
        """The [n] prefix must be stripped from text_raw."""
        data = _make_docx_numbered_with_runs([
            (1, [("contenu", {})]),
        ])
        p = tmp_path / "prefix_stripped.docx"
        p.write_bytes(data)

        from multicorpus_engine.importers.docx_numbered_lines import import_docx_numbered_lines
        report = import_docx_numbered_lines(db_conn, path=p, language="fr")

        row = db_conn.execute(
            "SELECT text_raw FROM units WHERE doc_id = ? AND unit_type = 'line'",
            (report.doc_id,),
        ).fetchone()
        assert "[1]" not in row["text_raw"], f"Prefix leaked into text_raw: {row['text_raw']!r}"


# ─── Tests: TEI round-trip ────────────────────────────────────────────────────

class TestTeiRoundTrip:
    def test_docx_italic_survives_tei_export(
        self, db_conn: sqlite3.Connection, tmp_path: Path
    ) -> None:
        """End-to-end: DOCX italic run → import → TEI export → <hi rend="italic"> in XML."""
        from multicorpus_engine.importers.docx_paragraphs import import_docx_paragraphs
        from multicorpus_engine.exporters.tei import export_tei

        data = _make_docx_with_runs([
            [("Voici ", {}), ("un mot", {"italic": True}), (" en italique.", {})],
        ])
        docx_path = tmp_path / "italic_roundtrip.docx"
        docx_path.write_bytes(data)

        report = import_docx_paragraphs(db_conn, path=docx_path, language="fr", title="Test")
        tei_path = tmp_path / "output.tei.xml"
        export_tei(db_conn, doc_id=report.doc_id, output_path=tei_path)

        root = ET.parse(str(tei_path)).getroot()
        body = root.find(f".//{_ns('body')}")
        assert body is not None
        p_elements = body.findall(f".//{_ns('p')}")
        assert p_elements, "No <p> in TEI body"
        p = p_elements[0]

        hi_elements = p.findall(_ns("hi"))
        assert len(hi_elements) == 1, f"Expected 1 <hi>, got {len(hi_elements)}"
        assert hi_elements[0].get("rend") == "italic"
        assert hi_elements[0].text == "un mot"
        assert p.text and "Voici" in p.text
        assert hi_elements[0].tail and "en italique" in hi_elements[0].tail

    def test_plain_docx_tei_export_no_hi(
        self, db_conn: sqlite3.Connection, tmp_path: Path
    ) -> None:
        """A DOCX with no styling must produce a TEI with no <hi> elements."""
        from multicorpus_engine.importers.docx_paragraphs import import_docx_paragraphs
        from multicorpus_engine.exporters.tei import export_tei

        data = _make_docx_with_runs([
            [("Paragraphe sans mise en forme.", {})],
        ])
        docx_path = tmp_path / "plain_roundtrip.docx"
        docx_path.write_bytes(data)

        report = import_docx_paragraphs(db_conn, path=docx_path, language="fr", title="Test")
        tei_path = tmp_path / "output_plain.tei.xml"
        export_tei(db_conn, doc_id=report.doc_id, output_path=tei_path)

        root = ET.parse(str(tei_path)).getroot()
        hi_elements = root.findall(f".//{_ns('hi')}")
        assert hi_elements == [], f"Unexpected <hi> elements: {hi_elements}"
