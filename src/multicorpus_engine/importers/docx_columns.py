"""Parcours d'un corps DOCX avec extraction par colonne de table.

Extrait de ``docx_numbered_lines`` (IMPO-01) pour être partagé avec
``docx_paragraphs`` : un bitexte en tableau à deux colonnes n'avait jusqu'ici
**aucun mode d'import valide**, ``column_index`` n'étant honoré que par le mode
« lignes numérotées », alors que ces documents ne portent aucun marqueur ``[n]``.

Le parcours est celui du mode numéroté, à l'identique — mêmes gardes sur les
fusions de cellules, mêmes compteurs. Les appelants comptent eux-mêmes ce qui les
regarde : le générateur dit d'où vient chaque paragraphe (premier niveau ou
cellule), parce que le mode numéroté ne pèse son ratio de lignes non numérotées
que sur le contenu **de la colonne**, jamais sur les paragraphes hors table.
"""

from __future__ import annotations

import logging
from dataclasses import dataclass
from typing import Iterator, Optional


@dataclass
class ColumnWalkStats:
    """Compteurs d'anomalies du parcours, tels que l'``ImportReport`` les publie."""

    tables_processed: int = 0
    rows_skipped_short: int = 0
    nested_tables_skipped: int = 0


def describe_tables(document) -> list[dict]:
    """Forme des tables de premier niveau, en ordre de lecture : ``{columns, rows}``.

    Sert à répondre à la seule question que l'écran d'import ne savait pas poser —
    *que contient ce fichier ?* — avant de demander une colonne. Sans elle,
    l'utilisateur devait deviner combien de colonnes existent, et le champ « colonne »
    n'avait de sens que pour qui connaissait déjà le document.

    Ne conclut rien : porter une table ne fait pas d'un document un bitexte. Mesuré le
    27 août 2026 sur le disque local, ``Conventions-Textes journalistiques`` porte
    **sept** tables de 5, 2, 2, 2, 2, 2 et 2 colonnes — de la mise en page, pas deux
    langues en regard. C'est à l'écran de montrer la forme, et à l'utilisateur de
    trancher sur pièces.

    Les tables **imbriquées** ne sont pas listées : le parcours d'extraction les saute
    aussi (en le signalant), les deux vues restent donc cohérentes.
    """
    from docx.table import Table as _DocxTable

    shapes: list[dict] = []
    for block in iter_body_blocks(document):
        if isinstance(block, _DocxTable):
            rows = len(block.rows)
            columns = len(block.rows[0].cells) if rows else 0
            shapes.append({"columns": columns, "rows": rows})
    return shapes


def column_walk_warnings(stats: ColumnWalkStats, column_index: int) -> list[str]:
    """Les pertes de données du parcours, rendues en avertissements d'``ImportReport``.

    Partagé par les deux modes DOCX : ces deux cas ne doivent rien à la numérotation,
    ils disent qu'une partie du document **n'a pas été importée**. Les libellés sont
    ceux du mode numéroté, mot pour mot — c'est le même fait, il se dit pareil.
    """
    warnings: list[str] = []
    if stats.rows_skipped_short > 0:
        warnings.append(
            f"{stats.rows_skipped_short} ligne(s) sur {stats.tables_processed} "
            f"table(s) ignorée(s) : colonne {column_index} absente (table plus "
            f"étroite ou cellule fusionnée venant d'une colonne précédente)."
        )
    if stats.nested_tables_skipped > 0:
        warnings.append(
            f"{stats.nested_tables_skipped} sous-table(s) imbriquée(s) ignorée(s) — "
            f"leur contenu n'a pas été importé."
        )
    return warnings


def iter_body_blocks(document):
    """Yield each top-level Paragraph or Table in document order.

    python-docx's ``Document.paragraphs`` skips paragraphs nested inside
    tables — we walk the body XML directly via lxml so column_index
    extraction can reach table content.
    """
    from docx.oxml.ns import qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    body = document.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield Table(child, document)


def is_vmerge_continuation(cell) -> bool:
    """True if *cell* is a vertical-merge CONTINUATION (not the start).

    In WordML, a vMerge has the start cell tagged ``<w:vMerge w:val="restart"/>``
    and continuation cells tagged ``<w:vMerge/>`` (absent val = continue).
    python-docx ``cell.merge`` sometimes leaves stale paragraph content
    in continuation cells, so simple ``id(cell)`` dedup is insufficient.
    We inspect the XML directly. Defensive : returns False on any error.
    """
    try:
        from docx.oxml.ns import qn
        tc_pr = cell._tc.tcPr
        if tc_pr is None:
            return False
        vmerge = tc_pr.find(qn("w:vMerge"))
        if vmerge is None:
            return False
        val = vmerge.get(qn("w:val"))
        return val != "restart"
    except Exception:
        return False


def iter_column_paragraphs(
    document,
    column_index: int,
    stats: ColumnWalkStats,
    log: Optional[logging.Logger] = None,
) -> Iterator[tuple[object, Optional[int], Optional[int]]]:
    """Parcourt le corps en ordre de lecture et plonge dans les tables à *column_index*.

    Rend ``(paragraphe, n_table, n_ligne)`` — ``n_table`` et ``n_ligne`` valent
    ``None`` pour un paragraphe de premier niveau, et sont 1-based sinon, de quoi
    reconstruire les journaux d'origine. *stats* est **muté** au fil du parcours.

    Les cas limites produisent des compteurs et des avertissements, jamais une
    perte silencieuse de données : ligne trop courte, fusion horizontale venant
    d'une colonne de gauche, fusion verticale de continuation, table imbriquée.
    """
    from docx.table import Table as _DocxTable
    from docx.text.paragraph import Paragraph as _DocxParagraph

    log = log or logging.getLogger(__name__)
    target_idx = column_index - 1

    for block in iter_body_blocks(document):
        if isinstance(block, _DocxParagraph):
            yield block, None, None
        elif isinstance(block, _DocxTable):
            stats.tables_processed += 1
            # Per-table dedup pour cellules fusionnées verticalement (vMerge) :
            # python-docx renvoie le MÊME élément <w:tc> pour les rows de
            # continuation d'un merge vertical. On garde les ÉLÉMENTS _tc
            # déjà vus (et non leur id()) : les proxies lxml sont
            # GC-ables et id() est réutilisé après collecte — d'où des
            # faux positifs de dedup en suite de tests complète. Conserver
            # la référence maintient le proxy en vie, donc la comparaison
            # `is` est stable.
            seen_target_tcs: list = []
            for row_idx, row in enumerate(block.rows):
                cells = row.cells
                if target_idx >= len(cells):
                    stats.rows_skipped_short += 1
                    continue
                target_cell = cells[target_idx]
                # Horizontal merge from a lower column: same Cell wrapper
                # appears at an earlier index. Skip rather que dupliquer
                # le contenu d'une cellule fusionnée venant de col 1.
                if target_idx > 0 and any(
                    cells[i] is target_cell for i in range(target_idx)
                ):
                    stats.rows_skipped_short += 1
                    continue
                # Vertical merge dedup : identité d'élément _tc via `is`.
                target_tc = target_cell._tc
                if any(target_tc is seen for seen in seen_target_tcs):
                    stats.rows_skipped_short += 1
                    continue
                # Défense secondaire : DOCX produits par Word (et non par
                # python-docx) marquent la continuation vMerge sans
                # val="restart" — le marqueur XML les attrape alors.
                if is_vmerge_continuation(target_cell):
                    stats.rows_skipped_short += 1
                    continue
                seen_target_tcs.append(target_tc)
                # Nested table inside the target cell — skip, warning.
                if target_cell.tables:
                    stats.nested_tables_skipped += len(target_cell.tables)
                    log.warning(
                        "Nested table in table %d row %d col %d — skipped",
                        stats.tables_processed, row_idx + 1, column_index,
                    )
                for para in target_cell.paragraphs:
                    yield para, stats.tables_processed, row_idx + 1
