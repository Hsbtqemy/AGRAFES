"""DOCX paragraphs importer.

Imports every non-empty paragraph as a line unit — no [n] numbering required.
The sequential paragraph position (1-based) is used as both n and external_id.

This makes all paragraphs alignable by position (monotone fallback alignment),
even when the source document has no explicit numbering convention.

See docs/DECISIONS.md ADR-012.
"""

from __future__ import annotations

import json
import logging
import sqlite3
from pathlib import Path
from typing import Optional

from ..unicode_policy import count_sep, normalize
from .docx_columns import (
    ColumnWalkStats,
    column_walk_warnings,
    describe_tables,
    iter_column_paragraphs,
)
from .docx_numbered_lines import ImportReport
from .import_guard import assert_not_duplicate_import, column_scoped_source_hash
from .parsed import ParsedDoc, ParsedUnit, file_sha256, insert_units
from .rich_text import para_to_rich_text

logger = logging.getLogger(__name__)


def parse_docx_paragraphs(
    path: str | Path,
    column_index: Optional[int] = None,
    run_logger: Optional[logging.Logger] = None,
) -> ParsedDoc:
    """Parse a DOCX into one line unit per non-empty paragraph, WITHOUT touching
    the DB. Shared by ``import_docx_paragraphs`` and ``/import/preview`` (A-02).
    Sequential position (1-based) is both n and external_id; Heading-styled
    paragraphs get unit_role="intertitre" + meta heading_level.

    ``column_index`` (1-based, IMPO-01) plonge dans les tables et n'y retient que
    la colonne demandée, paragraphes de premier niveau conservés en ordre de
    lecture. Sans lui, le comportement d'ADR-012 est inchangé : ``Document.paragraphs``
    ignore les tables, et un document qui n'est **qu'**un tableau rend zéro unité.
    C'est ce qui laissait un bitexte en tableau non numéroté sans aucun mode
    d'import valide — le mode numéroté le lisait, mais n'y trouvant pas de ``[n]``
    il en faisait un document entièrement ``structure``, donc hors index.
    """
    try:
        import docx  # python-docx
    except ImportError as exc:
        raise ImportError("python-docx is required: pip install python-docx") from exc

    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(f"DOCX file not found: {path}")

    _MAX_FILE_BYTES = 512 * 1024 * 1024  # 512 MiB
    if path.stat().st_size > _MAX_FILE_BYTES:
        raise ValueError(f"DOCX file too large (max {_MAX_FILE_BYTES // (1024 * 1024)} MiB)")

    # Validated here as in docx_numbered_lines, so the user gets a clear error
    # rather than a silent walk of column 0 (target_idx = -1 → la dernière colonne).
    if column_index is not None and column_index < 1:
        raise ValueError(f"column_index must be >= 1 (got {column_index})")

    log = run_logger or logger
    # IMPO-01 — l'identite d'un document extrait est (fichier, colonne).
    source_hash = column_scoped_source_hash(file_sha256(path), column_index)
    document = docx.Document(str(path))

    walk_stats = ColumnWalkStats()
    if column_index is None:
        paragraphs = document.paragraphs
    else:
        paragraphs = [
            para for para, _table_no, _row_no in iter_column_paragraphs(
                document, column_index, walk_stats, log
            )
        ]

    # First pass: collect paragraphs and detect headings.
    para_data: list[tuple[str, int | None]] = []
    for para in paragraphs:
        text_raw = para_to_rich_text(para)
        if not normalize(text_raw).strip():
            continue
        heading_level: int | None = None
        style_name = (para.style.name or "") if para.style else ""
        if style_name.startswith("Heading"):
            try:
                heading_level = int(style_name.split()[-1])
            except ValueError:
                heading_level = 1
        para_data.append((text_raw, heading_level))

    units: list[ParsedUnit] = []
    n = 0
    for text_raw, heading_level in para_data:
        n += 1
        text_norm = normalize(text_raw)
        sep_count = count_sep(text_raw)
        meta_dict: dict = {}
        if sep_count > 0:
            meta_dict["sep_count"] = sep_count
        if heading_level is not None:
            meta_dict["heading_level"] = heading_level
        meta = json.dumps(meta_dict) if meta_dict else None
        unit_role = "intertitre" if heading_level is not None else None
        units.append(ParsedUnit(
            n=n, unit_type="line", text_raw=text_raw, text_norm=text_norm,
            external_id=n, meta_json=meta, unit_role=unit_role,
        ))
        log.debug("Para n=%d type=line role=%s", n, unit_role)

    return ParsedDoc(
        units=units,
        doc_meta={},
        source_hash=source_hash,
        stats={
            "tables_processed": walk_stats.tables_processed,
            "rows_skipped_short": walk_stats.rows_skipped_short,
            "nested_tables_skipped": walk_stats.nested_tables_skipped,
            # Forme des tables, publiee MEME sans column_index : c'est ce qui permet
            # a l'ecran de dire ce que le fichier contient avant qu'on demande une
            # colonne (IMPO-01). Le document est deja ouvert, le cout est nul.
            "tables": describe_tables(document),
        },
    )


def import_docx_paragraphs(
    conn: sqlite3.Connection,
    path: str | Path,
    language: str,
    title: Optional[str] = None,
    doc_role: str = "standalone",
    resource_type: Optional[str] = None,
    column_index: Optional[int] = None,
    run_id: Optional[str] = None,
    run_logger: Optional[logging.Logger] = None,
    check_filename: bool = False,
    source_path: Optional[str] = None,
) -> ImportReport:
    """Import a DOCX file — every non-empty paragraph becomes a line unit.

    Unlike docx_numbered_lines, no [n] prefix is required. The sequential
    position (1-based) is stored as both n and external_id, so paragraphs
    are always monotone and gap-free.

    With ``column_index`` (IMPO-01), a two-column table bitext yields one document
    per column, both numbered 1..N by position — donc alignables par ancre dès
    l'import, sans qu'aucun marqueur ait à figurer dans la source.

    Returns an ImportReport. units_structure is always 0 (all units are lines).
    """
    path = Path(path)
    log = run_logger or logger
    log.info(
        "Starting import of %s (mode=docx_paragraphs, column_index=%s)", path, column_index
    )

    parsed = parse_docx_paragraphs(path, column_index=column_index, run_logger=run_logger)
    assert_not_duplicate_import(
        conn, path, parsed.source_hash,
        check_filename=check_filename, column_index=column_index,
    )
    source_hash = parsed.source_hash
    doc_title = title or path.stem
    utcnow = __import__("datetime").datetime.now(
        __import__("datetime").timezone.utc
    ).strftime("%Y-%m-%dT%H:%M:%SZ")

    has_headings = any(u.unit_role == "intertitre" for u in parsed.units)
    n_headings = sum(1 for u in parsed.units if u.unit_role == "intertitre")

    # Single transaction: document record + units
    try:
        cur = conn.execute(
            """
            INSERT INTO documents
                (title, language, doc_role, resource_type, meta_json, source_path, source_hash, created_at)
            VALUES (?, ?, ?, ?, NULL, ?, ?, ?)
            """,
            (doc_title, language, doc_role, resource_type, (source_path if source_path is not None else str(path)), source_hash, utcnow),
        )
        doc_id = cur.lastrowid
        log.info("Created document doc_id=%d title=%r", doc_id, doc_title)
        if has_headings:
            conn.execute(
                """
                INSERT OR IGNORE INTO unit_roles (name, label, color, icon, sort_order, category)
                VALUES ('intertitre', 'Intertitre', '#9333ea', '§', 0, 'structure')
                """
            )
        insert_units(conn, doc_id, parsed.units)
        conn.commit()
    except Exception:
        conn.rollback()
        raise

    report = ImportReport(
        doc_id=doc_id,
        units_total=len(parsed.units),
        units_line=len(parsed.units),
        units_structure=0,
        duplicates=[],
        holes=[],
        non_monotonic=[],
        tables_processed=parsed.stats.get("tables_processed", 0),
        rows_skipped_short=parsed.stats.get("rows_skipped_short", 0),
        nested_tables_skipped=parsed.stats.get("nested_tables_skipped", 0),
    )

    # IMPO-01 — une perte de donnees ne doit jamais etre muette : le mode numerote
    # avertissait deja, celui-ci comptait sans rien dire.
    if column_index is not None:
        for msg in column_walk_warnings(
            ColumnWalkStats(
                report.tables_processed,
                report.rows_skipped_short,
                report.nested_tables_skipped,
            ),
            column_index,
        ):
            report.warnings.append(msg)
            log.warning(msg)

    log.info("Import complete: %d paragraph units (%d headings)", report.units_total, n_headings)
    return report
