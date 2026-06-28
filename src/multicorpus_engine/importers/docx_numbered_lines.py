"""DOCX numbered-lines importer.

Reads a DOCX file where numbered paragraphs follow the pattern:
    [n] text content here

Rules (see docs/DECISIONS.md ADR-001, ADR-002, ADR-003):
- Paragraphs matching r'^\\[\\s*(\\d+)\\s*\\]\\s*(.+)$' → unit_type="line"
  - external_id = int(match.group(1))
  - text_raw = match.group(2) (prefix stripped, ¤ kept)
  - text_norm = normalize(text_raw)
- Non-matching paragraphs → unit_type="structure", external_id=NULL, NOT indexed
- Diagnostics: duplicates, holes, non-monotonic sequence
"""

from __future__ import annotations

import json
import logging
import re
import sqlite3
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

from ..unicode_policy import count_sep, normalize
from .import_guard import assert_not_duplicate_import
from .parsed import ParsedDoc, ParsedUnit, file_sha256, insert_units
from .rich_text import para_to_rich_text

_NUMBERED_RE = re.compile(r"^\[\s*(\d+)\s*\]\s*(.+)$", re.DOTALL)

# Si >50% des paragraphes d'une colonne demandée ne matchent pas `[N]`,
# emit a warning. Seuil minimum d'échantillon : 5 paragraphes pour éviter
# les faux positifs sur de petites tables. Constants figées ici.
COLUMN_UNNUMBERED_RATIO_THRESHOLD = 0.5
COLUMN_UNNUMBERED_MIN_SAMPLE = 5

logger = logging.getLogger(__name__)


@dataclass
class ImportReport:
    doc_id: int = 0
    units_total: int = 0
    units_line: int = 0
    units_structure: int = 0
    duplicates: list[int] = field(default_factory=list)
    holes: list[int] = field(default_factory=list)
    non_monotonic: list[int] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)
    # Statistiques d'extraction par table (renseignées seulement si column_index
    # est fourni à l'import — sinon valent 0).
    tables_processed: int = 0
    rows_skipped_short: int = 0
    nested_tables_skipped: int = 0

    def to_dict(self) -> dict:
        return {
            "doc_id": self.doc_id,
            "units_total": self.units_total,
            "units_line": self.units_line,
            "units_structure": self.units_structure,
            "duplicates": self.duplicates,
            "holes": self.holes,
            "non_monotonic": self.non_monotonic,
            "warnings": self.warnings,
            "tables_processed": self.tables_processed,
            "rows_skipped_short": self.rows_skipped_short,
            "nested_tables_skipped": self.nested_tables_skipped,
        }


def _paragraph_to_unit(
    para,
    n: int,
) -> tuple | None:
    """Convert one DOCX paragraph to a unit tuple, or None if blank.

    Returns ``(unit_type, n, ext_id_or_None, text_raw, text_norm, meta_or_None)``.
    Reused for both top-level paragraphs and cell paragraphs so the
    matching logic stays in one place.
    """
    rich = para_to_rich_text(para)
    plain = normalize(rich).strip()
    if not plain:
        return None
    m = _NUMBERED_RE.match(plain)
    if m:
        ext_id = int(m.group(1))
        # ENG-02: m.start(2) is an offset into `plain` (normalized + stripped);
        # slicing `rich` by it misaligns whenever normalize/strip changed the prefix
        # length. Re-match the marker against the lstripped rich text to strip exactly
        # the [n] prefix while preserving styling; fall back to the plain offset only
        # if the rich marker doesn't match (e.g. normalize rewrote the brackets).
        m_rich = _NUMBERED_RE.match(rich.lstrip())
        if m_rich:
            text_raw = m_rich.group(2)
        else:
            prefix_len = m.start(2)
            text_raw = rich[prefix_len:] if len(rich) >= prefix_len else m.group(2)
        text_norm = normalize(text_raw)
        sep_count = count_sep(text_raw)
        meta = json.dumps({"sep_count": sep_count}) if sep_count > 0 else None
        return ("line", n, ext_id, text_raw, text_norm, meta)
    text_raw = rich
    text_norm = normalize(text_raw)
    return ("structure", n, None, text_raw, text_norm, None)


def _iter_body_blocks(document):
    """Yield each top-level Paragraph or Table in document order.

    python-docx's ``Document.paragraphs`` skips paragraphs nested inside
    tables — we walk the body XML directly via lxml so column_index
    extraction can reach table content.
    """
    from docx.oxml.ns import qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    body = document.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield Table(child, document)


def _is_vmerge_continuation(cell) -> bool:
    """True if *cell* is a vertical-merge CONTINUATION (not the start).

    In WordML, a vMerge has the start cell tagged ``<w:vMerge w:val="restart"/>``
    and continuation cells tagged ``<w:vMerge/>`` (absent val = continue).
    python-docx ``cell.merge`` sometimes leaves stale paragraph content
    in continuation cells, so simple ``id(cell)`` dedup is insufficient.
    We inspect the XML directly. Defensive : returns False on any error.
    """
    try:
        from docx.oxml.ns import qn
        tc_pr = cell._tc.tcPr
        if tc_pr is None:
            return False
        vmerge = tc_pr.find(qn("w:vMerge"))
        if vmerge is None:
            return False
        val = vmerge.get(qn("w:val"))
        return val != "restart"
    except Exception:
        return False


def _analyze_external_ids(external_ids: list[int]) -> tuple[list[int], list[int], list[int]]:
    """Return (duplicates, holes, non_monotonic) from a sequence of external_ids."""
    seen: dict[int, int] = {}
    duplicates: list[int] = []
    non_monotonic: list[int] = []

    for i, eid in enumerate(external_ids):
        if eid in seen:
            if eid not in duplicates:
                duplicates.append(eid)
        seen[eid] = i
        if i > 0 and eid <= external_ids[i - 1]:
            non_monotonic.append(eid)

    # Holes: integers between min and max not present in the set
    unique = sorted(set(external_ids))
    holes: list[int] = []
    if unique:
        for expected in range(unique[0], unique[-1] + 1):
            if expected not in set(external_ids):
                holes.append(expected)

    return duplicates, holes, non_monotonic


def parse_docx_numbered_lines(
    path: str | Path,
    column_index: Optional[int] = None,
    run_logger: Optional[logging.Logger] = None,
) -> ParsedDoc:
    """Parse a DOCX (numbered-lines convention) into units WITHOUT touching the DB.

    Shared by ``import_docx_numbered_lines`` (write path) and the sidecar
    ``/import/preview`` so the parsing — including the column_index table walk and
    its vMerge dedup — lives in exactly one place (A-02). Raises
    ``ImportError`` / ``FileNotFoundError`` / ``ValueError`` like the importer did.
    """
    try:
        import docx  # python-docx
    except ImportError as exc:
        raise ImportError("python-docx is required: pip install python-docx") from exc

    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(f"DOCX file not found: {path}")

    _MAX_FILE_BYTES = 512 * 1024 * 1024  # 512 MiB
    if path.stat().st_size > _MAX_FILE_BYTES:
        raise ValueError(f"DOCX file too large (max {_MAX_FILE_BYTES // (1024 * 1024)} MiB)")

    log = run_logger or logger
    source_hash = file_sha256(path)

    # Validate column_index early so the user gets a clear error.
    if column_index is not None and column_index < 1:
        raise ValueError(f"column_index must be >= 1 (got {column_index})")

    document = docx.Document(str(path))

    # (unit_type, n, ext_id, text_raw, text_norm, meta) — doc_id added by the writer.
    units_parsed: list[tuple] = []

    # Per-table extraction counters (renseignés seulement quand column_index est set).
    tables_processed = 0
    rows_skipped_short = 0
    nested_tables_skipped = 0
    col_paragraphs_total = 0
    col_paragraphs_line = 0

    n = 0

    def _append_unit(unit: tuple) -> None:
        nonlocal n
        n += 1
        # Override the n placeholder produced by the helper with the live counter.
        units_parsed.append((unit[0], n, *unit[2:]))

    if column_index is None:
        # Legacy path — tables are skipped (python-docx Document.paragraphs
        # already ignores them). Unchanged behavior.
        for para in document.paragraphs:
            unit = _paragraph_to_unit(para, 0)  # n re-attributed by _append_unit
            if unit is None:
                continue
            _append_unit(unit)
            log.debug("Para n=%d type=%s", n, unit[0])
    else:
        # Column extraction — walk body in document order, dive into tables
        # at the requested column. Edge cases produce warnings + counters,
        # never silent data loss.
        from docx.table import Table as _DocxTable
        from docx.text.paragraph import Paragraph as _DocxParagraph
        target_idx = column_index - 1
        for block in _iter_body_blocks(document):
            if isinstance(block, _DocxParagraph):
                unit = _paragraph_to_unit(block, 0)
                if unit is None:
                    continue
                _append_unit(unit)
                log.debug("Top-level para n=%d type=%s", n, unit[0])
            elif isinstance(block, _DocxTable):
                tables_processed += 1
                # Per-table dedup pour cellules fusionnées verticalement (vMerge) :
                # python-docx renvoie le MÊME élément <w:tc> pour les rows de
                # continuation d'un merge vertical. On garde les ÉLÉMENTS _tc
                # déjà vus (et non leur id()) : les proxies lxml sont
                # GC-ables et id() est réutilisé après collecte — d'où des
                # faux positifs de dedup en suite de tests complète. Conserver
                # la référence maintient le proxy en vie, donc la comparaison
                # `is` est stable.
                seen_target_tcs: list = []
                for row_idx, row in enumerate(block.rows):
                    cells = row.cells
                    if target_idx >= len(cells):
                        rows_skipped_short += 1
                        continue
                    target_cell = cells[target_idx]
                    # Horizontal merge from a lower column: same Cell wrapper
                    # appears at an earlier index. Skip rather than dupliquer
                    # le contenu d'une cellule fusionnée venant de col 1.
                    if target_idx > 0 and any(
                        cells[i] is target_cell for i in range(target_idx)
                    ):
                        rows_skipped_short += 1
                        continue
                    # Vertical merge dedup : identité d'élément _tc via `is`.
                    target_tc = target_cell._tc
                    if any(target_tc is seen for seen in seen_target_tcs):
                        rows_skipped_short += 1
                        continue
                    # Défense secondaire : DOCX produits par Word (et non par
                    # python-docx) marquent la continuation vMerge sans
                    # val="restart" — le marqueur XML les attrape alors.
                    if _is_vmerge_continuation(target_cell):
                        rows_skipped_short += 1
                        continue
                    seen_target_tcs.append(target_tc)
                    # Nested table inside the target cell — skip, warning.
                    if target_cell.tables:
                        nested_tables_skipped += len(target_cell.tables)
                        log.warning(
                            "Nested table in table %d row %d col %d — skipped",
                            tables_processed, row_idx + 1, column_index,
                        )
                    for para in target_cell.paragraphs:
                        unit = _paragraph_to_unit(para, 0)
                        if unit is None:
                            continue
                        col_paragraphs_total += 1
                        if unit[0] == "line":
                            col_paragraphs_line += 1
                        _append_unit(unit)
                        log.debug(
                            "Table %d row %d col %d para n=%d type=%s",
                            tables_processed, row_idx + 1, column_index,
                            n, unit[0],
                        )

    units = [
        ParsedUnit(
            n=u[1], unit_type=u[0], text_raw=u[3], text_norm=u[4],
            external_id=u[2], meta_json=u[5],
        )
        for u in units_parsed
    ]
    return ParsedDoc(
        units=units,
        doc_meta={},
        source_hash=source_hash,
        stats={
            "tables_processed": tables_processed,
            "rows_skipped_short": rows_skipped_short,
            "nested_tables_skipped": nested_tables_skipped,
            "col_paragraphs_total": col_paragraphs_total,
            "col_paragraphs_line": col_paragraphs_line,
        },
    )


def import_docx_numbered_lines(
    conn: sqlite3.Connection,
    path: str | Path,
    language: str,
    title: Optional[str] = None,
    doc_role: str = "standalone",
    resource_type: Optional[str] = None,
    run_id: Optional[str] = None,
    run_logger: Optional[logging.Logger] = None,
    check_filename: bool = False,
    column_index: Optional[int] = None,
) -> ImportReport:
    """Import a DOCX file using the numbered-lines convention.

    Creates one row in `documents` and one row per paragraph in `units`.
    Returns an ImportReport with diagnostics.

    ``column_index`` controls table handling:
      - ``None`` (default) → tables are skipped entirely (legacy behavior).
      - ``>= 1`` → walk the body in document order; for each table, extract
        the cell at column ``column_index`` (1-based) and flatten its
        paragraphs. Pathological cases (table with fewer columns, merged
        cells coming from a lower column, nested sub-tables) are surfaced
        in ``ImportReport.warnings`` / new counters instead of failing
        silently.
    """
    path = Path(path)
    log = run_logger or logger
    log.info("Starting import of %s (mode=docx_numbered_lines)", path)

    parsed = parse_docx_numbered_lines(path, column_index=column_index, run_logger=run_logger)
    assert_not_duplicate_import(conn, path, parsed.source_hash, check_filename=check_filename)
    source_hash = parsed.source_hash
    doc_title = title or path.stem
    utcnow = __import__("datetime").datetime.now(
        __import__("datetime").timezone.utc
    ).strftime("%Y-%m-%dT%H:%M:%SZ")

    external_ids: list[int] = [
        u.external_id for u in parsed.units
        if u.unit_type == "line" and u.external_id is not None
    ]
    s = parsed.stats
    tables_processed = s["tables_processed"]
    rows_skipped_short = s["rows_skipped_short"]
    nested_tables_skipped = s["nested_tables_skipped"]
    col_paragraphs_total = s["col_paragraphs_total"]
    col_paragraphs_line = s["col_paragraphs_line"]

    # Single transaction: document record + units
    try:
        cur = conn.execute(
            """
            INSERT INTO documents
                (title, language, doc_role, resource_type, meta_json, source_path, source_hash, created_at)
            VALUES (?, ?, ?, ?, NULL, ?, ?, ?)
            """,
            (doc_title, language, doc_role, resource_type, str(path), source_hash, utcnow),
        )
        doc_id = cur.lastrowid
        log.info("Created document doc_id=%d title=%r", doc_id, doc_title)
        insert_units(conn, doc_id, parsed.units)
        conn.commit()
    except Exception:
        conn.rollback()
        raise

    # Build diagnostics
    duplicates, holes, non_monotonic = _analyze_external_ids(external_ids)

    report = ImportReport(
        doc_id=doc_id,
        units_total=len(parsed.units),
        units_line=len(external_ids),
        units_structure=len(parsed.units) - len(external_ids),
        duplicates=duplicates,
        holes=holes,
        non_monotonic=non_monotonic,
        tables_processed=tables_processed,
        rows_skipped_short=rows_skipped_short,
        nested_tables_skipped=nested_tables_skipped,
    )

    if duplicates:
        msg = f"Duplicate external_id(s) found: {duplicates}"
        report.warnings.append(msg)
        log.warning(msg)
    if holes:
        msg = f"Holes in external_id sequence: {holes}"
        report.warnings.append(msg)
        log.warning(msg)
    if non_monotonic:
        msg = f"Non-monotonic external_id(s): {non_monotonic}"
        report.warnings.append(msg)
        log.warning(msg)

    # Column-index warnings — transform "0 line units" silence into actionable signal.
    if column_index is not None:
        if rows_skipped_short > 0:
            msg = (
                f"{rows_skipped_short} ligne(s) sur {tables_processed} table(s) ignorée(s) : "
                f"colonne {column_index} absente (table plus étroite ou cellule fusionnée "
                f"venant d'une colonne précédente)."
            )
            report.warnings.append(msg)
            log.warning(msg)
        if nested_tables_skipped > 0:
            msg = (
                f"{nested_tables_skipped} sous-table(s) imbriquée(s) ignorée(s) — "
                f"leur contenu n'a pas été importé."
            )
            report.warnings.append(msg)
            log.warning(msg)
        if len(external_ids) == 0 and tables_processed > 0:
            msg = (
                f"0 unité ligne extraite de {tables_processed} table(s) à la colonne "
                f"{column_index}. Vérifier : la colonne contient-elle bien des "
                f"paragraphes numérotés `[N]` ?"
            )
            report.warnings.append(msg)
            log.warning(msg)
        elif (
            col_paragraphs_total >= COLUMN_UNNUMBERED_MIN_SAMPLE
            and (col_paragraphs_total - col_paragraphs_line) / col_paragraphs_total
                > COLUMN_UNNUMBERED_RATIO_THRESHOLD
        ):
            pct = round(
                100.0 * (col_paragraphs_total - col_paragraphs_line) / col_paragraphs_total
            )
            msg = (
                f"{pct}% des paragraphes de la colonne {column_index} ne portent pas "
                f"de numérotation `[N]`. Êtes-vous sûr d'avoir choisi la bonne colonne ?"
            )
            report.warnings.append(msg)
            log.warning(msg)

    log.info(
        "Import complete: %d units (%d line, %d structure)",
        report.units_total,
        report.units_line,
        report.units_structure,
    )
    return report
