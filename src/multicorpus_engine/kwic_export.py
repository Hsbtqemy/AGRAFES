"""Export KWIC concordance results to CSV, TXT, DOCX and ODT.

Each format receives a list of hit dicts with the following keys::

    doc_id        int
    doc_title     str
    unit_id       int
    unit_position int
    left          list[str]   — context tokens left of the match
    node          list[str]   — matched token(s)
    right         list[str]   — context tokens right of the match

Usage
-----
::

    from multicorpus_engine.kwic_export import export_kwic
    rows = export_kwic(hits, fmt="docx", out_path="/tmp/results.docx")
"""
from __future__ import annotations

import csv
from pathlib import Path


# ─── CSV ──────────────────────────────────────────────────────────────────────

def export_csv(hits: list[dict], out_path: str | Path) -> int:
    """Tab-separated CSV with BOM (Excel-friendly).

    Columns: doc_id, doc_title, unit_id, gauche, nœud, droite
    """
    path = Path(out_path)
    path.parent.mkdir(parents=True, exist_ok=True)
    with open(path, "w", newline="", encoding="utf-8-sig") as fh:
        w = csv.writer(fh, delimiter="\t")
        w.writerow(["doc_id", "doc_titre", "unit_id", "gauche", "nœud", "droite"])
        for h in hits:
            w.writerow([
                h["doc_id"],
                h.get("doc_title", ""),
                h["unit_id"],
                " ".join(h.get("left",  [])),
                " ".join(h.get("node",  [])),
                " ".join(h.get("right", [])),
            ])
    return len(hits)


# ─── TXT ──────────────────────────────────────────────────────────────────────

_LEFT_WIDTH  = 40
_RIGHT_WIDTH = 40


def export_txt(hits: list[dict], out_path: str | Path) -> int:
    """Fixed-width KWIC text file (classic concordancer layout).

    Each line: ``left_context  NODE  right_context  [doc#N u#M]``
    """
    path = Path(out_path)
    path.parent.mkdir(parents=True, exist_ok=True)
    with open(path, "w", encoding="utf-8") as fh:
        for h in hits:
            left  = " ".join(h.get("left",  []))
            node  = " ".join(h.get("node",  []))
            right = " ".join(h.get("right", []))
            ref   = f"[doc#{h['doc_id']} u#{h['unit_id']}]"
            fh.write(
                f"{left:>{_LEFT_WIDTH}}  {node}  {right:<{_RIGHT_WIDTH}}  {ref}\n"
            )
    return len(hits)


# ─── DOCX ─────────────────────────────────────────────────────────────────────

def export_docx(hits: list[dict], out_path: str | Path) -> int:
    """Word .docx table — node column in bold."""
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError as exc:
        raise RuntimeError(
            "python-docx is required for DOCX export. "
            "Install with: pip install python-docx"
        ) from exc

    doc = Document()
    doc.add_heading("KWIC — Résultats de concordance", level=1)

    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    for i, label in enumerate(["Doc", "Unité", "← Contexte", "Nœud", "Contexte →"]):
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(label)
        run.bold = True

    for h in hits:
        row = table.add_row().cells
        row[0].text = f"#{h['doc_id']}"
        row[1].text = str(h["unit_id"])
        row[2].text = " ".join(h.get("left",  []))
        # Node cell: bold
        p = row[3].paragraphs[0]
        run = p.add_run(" ".join(h.get("node", [])))
        run.bold = True
        row[4].text = " ".join(h.get("right", []))

    path = Path(out_path)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    return len(hits)


# ─── ODT ──────────────────────────────────────────────────────────────────────

def export_odt(hits: list[dict], out_path: str | Path) -> int:
    """LibreOffice ODT table — node column in bold."""
    try:
        from odf.opendocument import OpenDocumentText
        from odf.style import Style, TextProperties
        from odf.text import H, P, Span
        from odf.table import Table, TableColumn, TableRow, TableCell
    except ImportError as exc:
        raise RuntimeError(
            "odfpy is required for ODT export. "
            "Install with: pip install odfpy"
        ) from exc

    odtdoc = OpenDocumentText()

    bold_style = Style(name="Bold", family="text")
    bold_style.addElement(TextProperties(fontweight="bold"))
    odtdoc.styles.addElement(bold_style)

    heading = H(outlinelevel=1)
    heading.addText("KWIC — Résultats de concordance")
    odtdoc.text.addElement(heading)

    table = Table(name="KWIC")
    for _ in range(5):
        table.addElement(TableColumn())

    # Header row
    hdr_row = TableRow()
    for label in ["Doc", "Unité", "← Contexte", "Nœud", "Contexte →"]:
        cell = TableCell()
        p = P()
        span = Span(stylename=bold_style)
        span.addText(label)
        p.addElement(span)
        cell.addElement(p)
        hdr_row.addElement(cell)
    table.addElement(hdr_row)

    # Data rows
    for h in hits:
        row = TableRow()

        for text in [
            f"#{h['doc_id']}",
            str(h["unit_id"]),
            " ".join(h.get("left", [])),
        ]:
            cell = TableCell()
            p = P()
            p.addText(text)
            cell.addElement(p)
            row.addElement(cell)

        # Node cell (bold)
        cell = TableCell()
        p = P()
        span = Span(stylename=bold_style)
        span.addText(" ".join(h.get("node", [])))
        p.addElement(span)
        cell.addElement(p)
        row.addElement(cell)

        # Right context
        cell = TableCell()
        p = P()
        p.addText(" ".join(h.get("right", [])))
        cell.addElement(p)
        row.addElement(cell)

        table.addElement(row)

    odtdoc.text.addElement(table)

    path = Path(out_path)
    path.parent.mkdir(parents=True, exist_ok=True)
    odtdoc.save(str(path))
    return len(hits)


# ─── Dispatcher ───────────────────────────────────────────────────────────────

_FORMATS: dict[str, object] = {
    "csv":  export_csv,
    "txt":  export_txt,
    "docx": export_docx,
    "odt":  export_odt,
}

SUPPORTED_FORMATS = tuple(_FORMATS.keys())


def export_kwic(hits: list[dict], fmt: str, out_path: str | Path) -> int:
    """Export *hits* to *out_path* in *fmt* format.

    Parameters
    ----------
    hits:
        List of KWIC hit dicts (as returned by ``KwicHit.to_dict()``).
    fmt:
        One of ``"csv"``, ``"txt"``, ``"docx"``, ``"odt"``.
    out_path:
        Absolute path for the output file.

    Returns
    -------
    int
        Number of rows written.

    Raises
    ------
    ValueError
        If *fmt* is not recognised.
    RuntimeError
        If the required library for the format is not installed.
    """
    if fmt not in _FORMATS:
        raise ValueError(
            f"Unknown export format {fmt!r}. "
            f"Valid: {', '.join(SUPPORTED_FORMATS)}"
        )
    fn = _FORMATS[fmt]
    return fn(hits, out_path)  # type: ignore[operator]
