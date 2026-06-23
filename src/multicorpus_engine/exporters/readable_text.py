"""Readable text exporters (TXT / DOCX / ODT) for prepared corpus documents."""

from __future__ import annotations

import re
import sqlite3
import zipfile
from pathlib import Path
from typing import Any
from xml.sax.saxutils import escape


_VALID_FORMATS = {"txt", "docx", "odt"}
_VALID_SOURCE_FIELDS = {"text_norm", "text_raw"}

# ── ODT (OpenDocument Text) generation — stdlib only ──────────────────────────
# DOCX export uses python-docx; ODT has no project dependency (engine stays near
# stdlib), so we emit a minimal valid ODF 1.2 package by hand: a ZIP with an
# uncompressed `mimetype` first entry, a manifest, and `content.xml` carrying
# `text:h` (title) + `text:p` (lines) — the exact shape the ODT importer reads
# back (`importers/odt_common.py`), so exports round-trip.
_ODT_MIMETYPE = "application/vnd.oasis.opendocument.text"

_ODT_MANIFEST = (
    '<?xml version="1.0" encoding="UTF-8"?>\n'
    '<manifest:manifest '
    'xmlns:manifest="urn:oasis:names:tc:opendocument:xmlns:manifest:1.0" '
    'manifest:version="1.2">'
    f'<manifest:file-entry manifest:full-path="/" manifest:media-type="{_ODT_MIMETYPE}"/>'
    '<manifest:file-entry manifest:full-path="content.xml" manifest:media-type="text/xml"/>'
    '</manifest:manifest>'
)


def _odt_content_xml(heading: str, lines: list[str]) -> str:
    """Build content.xml: heading as `text:h`, each line as a `text:p`."""
    body = [f'<text:h text:outline-level="1">{escape(heading)}</text:h>']
    body.extend(f"<text:p>{escape(line)}</text:p>" for line in lines)
    return (
        '<?xml version="1.0" encoding="UTF-8"?>\n'
        '<office:document-content '
        'xmlns:office="urn:oasis:names:tc:opendocument:xmlns:office:1.0" '
        'xmlns:text="urn:oasis:names:tc:opendocument:xmlns:text:1.0" '
        'office:version="1.2"><office:body><office:text>'
        + "".join(body)
        + "</office:text></office:body></office:document-content>"
    )


def _write_odt(dest: Path, heading: str, lines: list[str]) -> None:
    """Write a minimal valid .odt package (ZIP) for one document."""
    with zipfile.ZipFile(dest, "w", zipfile.ZIP_DEFLATED) as zf:
        # `mimetype` must be the first entry and STORED (uncompressed) per ODF spec.
        info = zipfile.ZipInfo("mimetype")
        info.compress_type = zipfile.ZIP_STORED
        zf.writestr(info, _ODT_MIMETYPE)
        zf.writestr("META-INF/manifest.xml", _ODT_MANIFEST)
        zf.writestr("content.xml", _odt_content_xml(heading, lines))


def _slugify(text: str) -> str:
    text = (text or "").strip().lower()
    text = re.sub(r"[^a-z0-9._-]+", "_", text)
    return text.strip("._-") or "document"


def _resolve_doc_ids(conn: sqlite3.Connection, doc_ids: list[int] | None) -> list[int]:
    if doc_ids is None:
        rows = conn.execute("SELECT doc_id FROM documents ORDER BY doc_id").fetchall()
        return [int(r[0]) for r in rows]
    return [int(d) for d in doc_ids]


def _load_doc_meta(conn: sqlite3.Connection, doc_id: int) -> tuple[str, str]:
    row = conn.execute(
        "SELECT title, language FROM documents WHERE doc_id = ?",
        (doc_id,),
    ).fetchone()
    if row is None:
        raise ValueError(f"Unknown doc_id: {doc_id}")
    title = str(row[0] or f"doc_{doc_id}")
    language = str(row[1] or "und")
    return title, language


def _load_doc_units(
    conn: sqlite3.Connection,
    *,
    doc_id: int,
    include_structure: bool,
) -> list[sqlite3.Row]:
    if include_structure:
        rows = conn.execute(
            """
            SELECT unit_type, n, external_id, text_raw, text_norm
            FROM units
            WHERE doc_id = ?
            ORDER BY n
            """,
            (doc_id,),
        ).fetchall()
    else:
        rows = conn.execute(
            """
            SELECT unit_type, n, external_id, text_raw, text_norm
            FROM units
            WHERE doc_id = ? AND unit_type = 'line'
            ORDER BY n
            """,
            (doc_id,),
        ).fetchall()
    return list(rows)


def _render_unit_text(
    row: sqlite3.Row,
    *,
    source_field: str,
    include_external_id: bool,
) -> str:
    text = row[source_field] or ""
    if include_external_id and row["external_id"] is not None:
        return f"[{int(row['external_id']):04d}] {text}"
    return str(text)


def export_readable_text(
    conn: sqlite3.Connection,
    *,
    out_dir: str | Path,
    doc_ids: list[int] | None = None,
    fmt: str = "txt",
    include_structure: bool = False,
    include_external_id: bool = True,
    source_field: str = "text_norm",
) -> dict[str, Any]:
    """Export selected documents as readable TXT, DOCX or ODT files.

    Args:
        conn: SQLite connection.
        out_dir: Destination directory.
        doc_ids: Selected document IDs (None -> all docs).
        fmt: "txt", "docx" or "odt".
        include_structure: Include `unit_type='structure'` rows when true.
        include_external_id: Prefix line units with `[0001]` style anchors.
        source_field: `text_norm` (default) or `text_raw`.
    """

    fmt_norm = str(fmt or "txt").strip().lower()
    if fmt_norm not in _VALID_FORMATS:
        raise ValueError(f"Unsupported readable text format: {fmt!r}")
    if source_field not in _VALID_SOURCE_FIELDS:
        raise ValueError("source_field must be 'text_norm' or 'text_raw'")

    out_path = Path(out_dir)
    out_path.mkdir(parents=True, exist_ok=True)

    ids = _resolve_doc_ids(conn, doc_ids)
    files_created: list[str] = []

    for doc_id in ids:
        title, language = _load_doc_meta(conn, doc_id)
        safe_title = _slugify(title)
        dest = out_path / f"doc_{doc_id}_{safe_title}.{fmt_norm}"
        rows = _load_doc_units(conn, doc_id=doc_id, include_structure=include_structure)

        rendered_lines = [
            _render_unit_text(
                row,
                source_field=source_field,
                include_external_id=include_external_id,
            )
            for row in rows
        ]

        if fmt_norm == "txt":
            with dest.open("w", encoding="utf-8", newline="\n") as f:
                f.write(f"# {title} [{language}]\n\n")
                for line in rendered_lines:
                    f.write(line)
                    f.write("\n")
        elif fmt_norm == "odt":
            _write_odt(dest, f"{title} [{language}]", rendered_lines)
        else:
            try:
                import docx  # python-docx
            except ImportError as exc:  # pragma: no cover - dependency is required in project deps
                raise ImportError("python-docx is required for DOCX export") from exc
            document = docx.Document()
            document.add_heading(f"{title} [{language}]", level=1)
            for line in rendered_lines:
                document.add_paragraph(line)
            document.save(dest)

        files_created.append(str(dest))

    return {
        "out_dir": str(out_path),
        "count": len(files_created),
        "format": fmt_norm,
        "files_created": files_created,
    }

